import io
from docx import Document as DocxDocument
from PIL import Image
from reviewer.parsing.docx_reader import read_docx


def _png_bytes() -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (8, 8), "white").save(buf, format="PNG")
    return buf.getvalue()


def _make_docx() -> bytes:
    d = DocxDocument()
    d.add_paragraph("Intro paragraph.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Cell A"
    table.rows[0].cells[1].text = "Cell B"
    d.add_picture(io.BytesIO(_png_bytes()))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_read_docx_extracts_paragraphs_tables_and_image_ocr():
    pc = read_docx(_make_docx(), ocr=lambda b, m: "IMAGE_TEXT")
    assert "Intro paragraph." in pc.text
    assert "Cell A" in pc.text and "Cell B" in pc.text
    assert "IMAGE_TEXT" in pc.text


def test_read_docx_swallows_ocr_failure_and_keeps_text():
    # A failing OCR on an embedded image must not abort extraction of the rest.
    def boom(b, m):
        raise RuntimeError("ocr down")

    pc = read_docx(_make_docx(), ocr=boom)
    assert "Intro paragraph." in pc.text
    assert "Cell A" in pc.text and "Cell B" in pc.text
