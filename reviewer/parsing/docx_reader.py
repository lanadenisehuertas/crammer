import io

from docx import Document as DocxDocument

from reviewer.parsing.base import OcrFn, ParsedContent


def read_docx(data: bytes, ocr: OcrFn) -> ParsedContent:
    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/image"):
            try:
                blob = rel.target_part.blob
                media_type = rel.target_part.content_type
                transcription = ocr(blob, media_type)
            except Exception:
                transcription = ""
            if transcription.strip():
                parts.append(transcription.strip())

    return ParsedContent(text="\n\n".join(parts))
